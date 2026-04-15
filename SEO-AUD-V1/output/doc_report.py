import json
import os
from typing import Any, Dict, List
from urllib.parse import urlparse

from docx import Document

from utils.helpers import safe_filename, utc_now_iso


def _table_from_rows(document: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Light List Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, value in enumerate(rows[0]):
        hdr_cells[idx].text = value

    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def run(final_audit: Dict[str, Any], output_dir: str, enable_google_docs: bool = False) -> Dict[str, Any]:
    os.makedirs(output_dir, exist_ok=True)

    url = final_audit.get("url", "unknown")
    host = urlparse(url).netloc or "unknown-host"
    timestamp = utc_now_iso().replace(":", "").replace("+00:00", "Z")

    doc_name = f"{safe_filename(host)}_{safe_filename(timestamp)}.docx"
    doc_path = os.path.join(output_dir, doc_name)

    analysis = final_audit.get("analysis", {})
    ai = final_audit.get("ai", {})
    summary = analysis.get("summary", {})
    priorities = analysis.get("priority_findings", [])
    ai_report = ai.get("report", {})

    document = Document()
    document.add_heading("SEO Audit Report", level=1)
    document.add_paragraph(f"URL: {url}")
    document.add_paragraph(f"Generated at (UTC): {final_audit.get('timestamp_utc', '')}")

    document.add_heading("Summary", level=2)
    document.add_paragraph(json.dumps(summary, indent=2, ensure_ascii=True))

    document.add_heading("Priority Findings", level=2)
    if priorities:
        rows = [["Priority", "Check", "Status", "Severity", "Message"]]
        for idx, item in enumerate(priorities[:12], start=1):
            rows.append(
                [
                    str(idx),
                    str(item.get("check_id", "")),
                    str(item.get("status", "")),
                    str(item.get("severity", "")),
                    str(item.get("message", "")),
                ]
            )
        _table_from_rows(document, rows)
    else:
        document.add_paragraph("No priority findings available.")

    document.add_heading("AI Narrative", level=2)
    text = ai_report.get("text", "AI report is not available for this run.")
    document.add_paragraph(text)

    document.save(doc_path)

    google_docs_payload = {
        "enabled": bool(enable_google_docs),
        "status": "not_requested" if not enable_google_docs else "pending_external_automation",
        "note": "Integrate with Google Docs API or automation worker to publish this DOCX content.",
    }

    return {
        "module": "output.doc_report",
        "status": "ok",
        "url": url,
        "path": doc_path,
        "google_docs": google_docs_payload,
    }
